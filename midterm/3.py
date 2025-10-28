import streamlit as st
from openai import OpenAI
import os
from dotenv import load_dotenv
import PyPDF2
from docx import Document
from docx.shared import Inches
from datetime import datetime
import io
from pathlib import Path

# .env 파일 로드
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# Streamlit 페이지 설정
st.set_page_config(
    page_title="문서 요약 애플리케이션",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS 스타일링
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        color: #2E8B57;
        text-align: center;
        margin-bottom: 2rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
    }
    .summary-container {
        background-color: #f0f8ff;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 5px solid #2E8B57;
    }
    .text-container {
        background-color: #f9f9f9;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        max-height: 300px;
        overflow-y: auto;
        border: 1px solid #ddd;
    }
    .info-badge {
        background-color: #17a2b8;
        color: white;
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        font-size: 0.8rem;
        margin: 0.2rem;
        display: inline-block;
    }
    .success-badge {
        background-color: #28a745;
        color: white;
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        font-size: 0.8rem;
        margin: 0.2rem;
        display: inline-block;
    }
</style>
""", unsafe_allow_html=True)

# 메인 헤더
st.markdown('<h1 class="main-header">📄 문서 요약 애플리케이션</h1>', unsafe_allow_html=True)

# 사이드바 설정
with st.sidebar:
    st.header("⚙️ 설정")
    
    # 요약 길이 설정
    summary_length = st.selectbox(
        "요약 길이",
        ["짧게", "보통", "자세히"],
        index=1
    )
    
    # 요약 언어 설정
    summary_language = st.selectbox(
        "요약 언어",
        ["한국어", "English"],
        index=0
    )
    
    # 최대 토큰 수 설정
    max_tokens = st.slider("최대 응답 길이", 100, 2000, 800)
    
    st.divider()
    
    # 앱 정보
    st.subheader("📋 기능")
    st.markdown('<span class="info-badge">PDF 텍스트 추출</span>', unsafe_allow_html=True)
    st.markdown('<span class="info-badge">AI 기반 요약</span>', unsafe_allow_html=True)
    st.markdown('<span class="info-badge">DOCX 파일 저장</span>', unsafe_allow_html=True)
    st.markdown('<span class="info-badge">다국어 지원</span>', unsafe_allow_html=True)

# PDF에서 텍스트 추출 함수
def extract_text_from_pdf(pdf_file):
    """PDF 파일에서 텍스트를 추출합니다."""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text + "\n"
        
        return text.strip()
    except Exception as e:
        st.error(f"PDF 읽기 오류: {str(e)}")
        return None

# GPT를 사용한 텍스트 요약 함수
def summarize_text(text, length, language):
    """GPT를 사용하여 텍스트를 요약합니다."""
    try:
        # 요약 길이에 따른 프롬프트 설정
        length_instructions = {
            "짧게": "3-4문장으로 핵심만 간단하게",
            "보통": "1-2개 단락으로 적당한 길이로",
            "자세히": "3-4개 단락으로 상세하게"
        }
        
        # 언어에 따른 프롬프트 설정
        if language == "한국어":
            system_prompt = f"""당신은 전문적인 문서 요약 전문가입니다. 
            주어진 텍스트를 {length_instructions[length]} 한국어로 요약해주세요.
            
            요약 시 다음 사항을 고려해주세요:
            - 핵심 내용과 주요 포인트 포함
            - 논리적인 구조로 정리
            - 명확하고 이해하기 쉬운 문장으로 작성
            - 원문의 의도와 맥락 유지"""
        else:
            system_prompt = f"""You are a professional document summarization expert. 
            Please summarize the given text {length_instructions[length]} in English.
            
            Please consider the following when summarizing:
            - Include key content and main points
            - Organize with logical structure
            - Write in clear and easy-to-understand sentences
            - Maintain the intent and context of the original text"""
        
        # OpenAI API 호출
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"다음 텍스트를 요약해주세요:\n\n{text}"}
            ],
            max_tokens=max_tokens,
            temperature=0.7
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.error(f"요약 생성 오류: {str(e)}")
        return None

# DOCX 파일 생성 함수
def create_docx_file(original_text, summary_text, filename):
    """요약 내용을 DOCX 파일로 바탕화면에 저장합니다."""
    try:
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading('문서 요약 보고서', 0)
        title.alignment = 1  # 가운데 정렬
        
        # 생성 정보 추가
        doc.add_paragraph(f"생성 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}")
        doc.add_paragraph(f"원본 파일: {filename}")
        doc.add_paragraph("=" * 50)
        
        # 요약 내용 추가
        doc.add_heading('📋 요약 내용', level=1)
        summary_paragraph = doc.add_paragraph(summary_text)
        
        # 원본 텍스트 정보 추가
        doc.add_heading('📄 원본 텍스트 정보', level=1)
        doc.add_paragraph(f"원본 텍스트 길이: {len(original_text):,} 자")
        
        # 원본 텍스트 일부 추가 (너무 길면 처음 500자만)
        doc.add_heading('📖 원본 텍스트 미리보기', level=2)
        preview_text = original_text[:500] + "..." if len(original_text) > 500 else original_text
        doc.add_paragraph(preview_text)
        
        # 🔥 핵심 변경: 바탕화면 경로에 직접 저장
        desktop_path = Path.home() / "Desktop"
        
        # 파일명 생성 (원본 파일명 기반)
        base_name = Path(filename).stem if filename else "문서요약"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"{base_name}_요약_{timestamp}.docx"
        
        # 바탕화면에 저장할 전체 경로
        save_path = desktop_path / docx_filename
        
        # 바탕화면에 파일 저장
        doc.save(str(save_path))
        
        # 메모리 스트림도 반환 (다운로드 버튼용)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream, str(save_path)
        
    except Exception as e:
        st.error(f"DOCX 파일 생성 오류: {str(e)}")
        return None, None

# 세션 상태 초기화
if "extracted_text" not in st.session_state:
    st.session_state.extracted_text = ""
if "summary_result" not in st.session_state:
    st.session_state.summary_result = ""
if "filename" not in st.session_state:
    st.session_state.filename = ""

# 메인 인터페이스
st.subheader("📤 PDF 파일 업로드")

uploaded_file = st.file_uploader(
    "PDF 파일을 선택하세요",
    type=['pdf'],
    help="PDF 파일만 업로드 가능합니다."
)

if uploaded_file is not None:
    st.session_state.filename = uploaded_file.name
    
    # 파일 정보 표시
    file_details = {
        "파일명": uploaded_file.name,
        "파일 크기": f"{uploaded_file.size / 1024:.2f} KB"
    }
    
    col1, col2 = st.columns(2)
    with col1:
        st.write("**파일 정보:**")
        for key, value in file_details.items():
            st.write(f"- {key}: {value}")
    
    with col2:
        if st.button("📖 텍스트 추출", type="primary"):
            with st.spinner("PDF에서 텍스트를 추출하고 있습니다..."):
                extracted_text = extract_text_from_pdf(uploaded_file)
                
                if extracted_text:
                    st.session_state.extracted_text = extracted_text
                    st.success("✅ 텍스트 추출이 완료되었습니다!")
                    st.markdown('<span class="success-badge">추출 완료</span>', unsafe_allow_html=True)

# 추출된 텍스트 표시 및 처리
if st.session_state.extracted_text:
    text_length = len(st.session_state.extracted_text)
    
    st.subheader("📄 추출된 텍스트")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("텍스트 길이", f"{text_length:,} 자")
    with col2:
        st.metric("단어 수 (근사)", f"{len(st.session_state.extracted_text.split()):,} 개")
    with col3:
        pages_estimate = max(1, text_length // 2000)
        st.metric("예상 페이지", f"{pages_estimate} 페이지")
    
    # 텍스트 미리보기
    with st.expander("📖 원본 텍스트 미리보기"):
        st.markdown('<div class="text-container">', unsafe_allow_html=True)
        preview_text = st.session_state.extracted_text[:1000] + "..." if len(st.session_state.extracted_text) > 1000 else st.session_state.extracted_text
        st.text(preview_text)
        st.markdown('</div>', unsafe_allow_html=True)
    
    # 요약 처리
    st.subheader("🎯 문서 요약")
    
    if st.button("📝 요약 생성", type="primary"):
        if text_length < 1000:
            # 1000자 미만인 경우
            st.session_state.summary_result = "요청한 문서가 짧아 요약 대신 원본 텍스트를 출력합니다."
            st.warning("⚠️ " + st.session_state.summary_result)
            
            st.markdown('<div class="summary-container">', unsafe_allow_html=True)
            st.markdown("**원본 텍스트:**")
            st.text(st.session_state.extracted_text)
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            # 1000자 이상인 경우 GPT API 호출
            with st.spinner("AI가 문서를 요약하고 있습니다..."):
                summary = summarize_text(
                    st.session_state.extracted_text, 
                    summary_length, 
                    summary_language
                )
                
                if summary:
                    st.session_state.summary_result = summary
                    st.success("✅ 요약이 완료되었습니다!")
                    
                    st.markdown('<div class="summary-container">', unsafe_allow_html=True)
                    st.markdown("**📋 요약 결과:**")
                    st.write(summary)
                    st.markdown('</div>', unsafe_allow_html=True)
    
    # DOCX 파일 다운로드 및 바탕화면 저장
    if st.session_state.summary_result:
        st.subheader("💾 파일 저장 및 다운로드")
        
        # 파일명 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"요약_{st.session_state.filename.replace('.pdf', '')}_{timestamp}.docx"
        
        # DOCX 파일 생성 (바탕화면 저장 + 메모리 스트림)
        result = create_docx_file(
            st.session_state.extracted_text,
            st.session_state.summary_result,
            st.session_state.filename
        )
        
        if result[0] is not None:  # 파일 생성 성공
            docx_file, desktop_path = result
            
            # 🔥 바탕화면 저장 성공 메시지
            st.success(f"✅ 파일이 바탕화면에 자동 저장되었습니다!")
            st.info(f"📁 저장 위치: {desktop_path}")
            
            # 추가 다운로드 옵션도 제공
            st.download_button(
                label="📄 추가 다운로드 (브라우저)",
                data=docx_file.getvalue(),
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="secondary",
                help="바탕화면 저장 외에 추가로 다운로드하려면 클릭하세요"
            )

# 하단 정보
st.divider()
with st.expander("ℹ️ 사용 안내"):
    st.write("""
    **📖 사용 방법:**
    1. PDF 파일을 업로드합니다
    2. '텍스트 추출' 버튼을 클릭합니다
    3. 추출된 텍스트를 확인합니다
    4. '요약 생성' 버튼을 클릭합니다
    5. 📁 **요약 파일이 자동으로 바탕화면에 저장됩니다!**
    6. 필요시 '추가 다운로드' 버튼으로 브라우저 다운로드도 가능합니다
    
    **✨ 새로운 기능:**
    - 🖥️ 바탕화면 자동 저장: 별도 경로 선택 없이 바로 사용 가능
    - 📝 파일명 자동 생성: 원본파일명_요약_날짜시간.docx
    
    **⚠️ 주의사항:**
    - 1000자 미만의 문서는 요약하지 않고 원본을 표시합니다
    - PDF 파일의 이미지나 복잡한 레이아웃은 추출되지 않을 수 있습니다
    - 생성된 DOCX 파일에는 요약과 원본 정보가 모두 포함됩니다
    """)

# 푸터
st.markdown("---")
st.markdown("📄 **문서 요약 애플리케이션** | Powered by OpenAI GPT")
